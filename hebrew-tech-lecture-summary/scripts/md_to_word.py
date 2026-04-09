"""
Convert a Markdown summary file to a formatted Word (.docx) document.
Handles: headings (H1-H3), bold, italic, bold+italic, blockquotes,
         horizontal rules, bullet lists, metadata lines (**Key:** Value),
         and inline code. RTL (Hebrew) by default.

Usage:
    python md_to_word.py <input.md> [output.docx]
"""

import sys
import os
import re

try:
    import docx
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ── RTL helpers ────────────────────────────────────────────────────────────────

def _set_para_rtl(para, is_rtl):
    """Set paragraph direction and alignment."""
    para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT if is_rtl else WD_PARAGRAPH_ALIGNMENT.LEFT
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1' if is_rtl else '0')
    pPr.append(bidi)


def _set_run_rtl(run, is_rtl):
    """Mark a run as RTL."""
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1' if is_rtl else '0')
    rPr.append(rtl)


# ── Inline Markdown → runs ─────────────────────────────────────────────────────

# Matches: ***bold+italic***, **bold**, *italic*, `code`
INLINE_RE = re.compile(
    r'(\*\*\*(.+?)\*\*\*'
    r'|\*\*(.+?)\*\*'
    r'|\*(.+?)\*'
    r'|`(.+?)`)',
    re.DOTALL
)


def _add_inline(para, text, is_rtl, base_size=None, base_color=None, base_bold=False, base_italic=False):
    """Parse inline markdown and add styled runs to para."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        # plain text before match
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
            run.bold = base_bold
            run.italic = base_italic
            if base_size:
                run.font.size = Pt(base_size)
            if base_color:
                run.font.color.rgb = base_color
            _set_run_rtl(run, is_rtl)

        segment = m.group(0)
        if segment.startswith('***'):
            run = para.add_run(m.group(2))
            run.bold = True
            run.italic = True
        elif segment.startswith('**'):
            run = para.add_run(m.group(3))
            run.bold = True
            run.italic = base_italic
        elif segment.startswith('*'):
            run = para.add_run(m.group(4))
            run.bold = base_bold
            run.italic = True
        else:  # backtick code
            run = para.add_run(m.group(5))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

        if base_size and not segment.startswith('`'):
            run.font.size = Pt(base_size)
        if base_color:
            run.font.color.rgb = base_color
        _set_run_rtl(run, is_rtl)
        pos = m.end()

    # trailing plain text
    if pos < len(text):
        run = para.add_run(text[pos:])
        run.bold = base_bold
        run.italic = base_italic
        if base_size:
            run.font.size = Pt(base_size)
        if base_color:
            run.font.color.rgb = base_color
        _set_run_rtl(run, is_rtl)


# ── Metadata line (**Key:** value) ─────────────────────────────────────────────

META_RE = re.compile(r'^\*\*(.+?):\*\*\s*(.*)')


def _add_meta_line(doc, line, is_rtl):
    m = META_RE.match(line)
    para = doc.add_paragraph()
    _set_para_rtl(para, is_rtl)
    para.paragraph_format.space_after = Pt(2)
    if m:
        key_run = para.add_run(m.group(1) + ': ')
        key_run.bold = True
        key_run.font.size = Pt(11)
        _set_run_rtl(key_run, is_rtl)
        _add_inline(para, m.group(2), is_rtl, base_size=11)
    else:
        _add_inline(para, line, is_rtl, base_size=11)
    return para


# ── Horizontal rule ────────────────────────────────────────────────────────────

def _add_hr(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)


# ── Blockquote ─────────────────────────────────────────────────────────────────

QUOTE_COLOR = RGBColor(0x44, 0x44, 0x44)


def _add_blockquote(doc, text, is_rtl):
    para = doc.add_paragraph()
    _set_para_rtl(para, is_rtl)
    para.paragraph_format.left_indent = Cm(1)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    side = OxmlElement('w:right' if is_rtl else 'w:left')
    side.set(qn('w:val'), 'single')
    side.set(qn('w:sz'), '18')
    side.set(qn('w:space'), '10')
    side.set(qn('w:color'), '888888')
    pBdr.append(side)
    pPr.append(pBdr)
    _add_inline(para, text, is_rtl, base_size=11, base_color=QUOTE_COLOR, base_italic=True)


# ── Heading sizes ──────────────────────────────────────────────────────────────

HEADING_SIZE = {1: 20, 2: 16, 3: 13}
HEADING_COLOR = {
    1: RGBColor(0x1F, 0x35, 0x64),
    2: RGBColor(0x2E, 0x74, 0xB5),
    3: RGBColor(0x1F, 0x35, 0x64),
}


def _add_heading(doc, text, level, is_rtl):
    para = doc.add_paragraph()
    _set_para_rtl(para, is_rtl)
    size = HEADING_SIZE.get(level, 12)
    color = HEADING_COLOR.get(level)
    para.paragraph_format.space_before = Pt(14 if level == 1 else 10 if level == 2 else 6)
    para.paragraph_format.space_after = Pt(4)
    _add_inline(para, text, is_rtl, base_size=size, base_color=color, base_bold=True)


# ── Bullet list ────────────────────────────────────────────────────────────────

def _add_bullet(doc, text, is_rtl):
    para = doc.add_paragraph(style='List Bullet')
    _set_para_rtl(para, is_rtl)
    para.paragraph_format.space_after = Pt(2)
    _add_inline(para, text, is_rtl, base_size=11)


# ── Code block ─────────────────────────────────────────────────────────────────

def _add_code_line(doc, text, is_rtl):
    para = doc.add_paragraph()
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # code is always LTR
    para.paragraph_format.left_indent = Cm(0.5)
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)


# ── Main converter ─────────────────────────────────────────────────────────────

def md_to_word(input_file, output_file, is_rtl=True):
    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.read().splitlines()

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # Default body font
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    in_code_block = False
    i = 0
    while i < len(lines):
        line = lines[i]

        # ── Code fence ──────────────────────────────────────────────────────
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            _add_code_line(doc, line, is_rtl)
            i += 1
            continue

        # ── Horizontal rule ─────────────────────────────────────────────────
        if re.match(r'^---+$', line.strip()):
            _add_hr(doc)
            i += 1
            continue

        # ── Headings ────────────────────────────────────────────────────────
        heading_match = re.match(r'^(#{1,3})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            _add_heading(doc, text, level, is_rtl)
            i += 1
            continue

        # ── Blockquote ──────────────────────────────────────────────────────
        if line.startswith('>'):
            text = re.sub(r'^>\s?', '', line)
            _add_blockquote(doc, text, is_rtl)
            i += 1
            continue

        # ── Bullet list ─────────────────────────────────────────────────────
        bullet_match = re.match(r'^[-*]\s+(.*)', line)
        if bullet_match:
            _add_bullet(doc, bullet_match.group(1), is_rtl)
            i += 1
            continue

        # ── Empty line ──────────────────────────────────────────────────────
        if line.strip() == '':
            # small spacer paragraph
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # ── Metadata line (**Key:** value) ──────────────────────────────────
        if META_RE.match(line):
            _add_meta_line(doc, line, is_rtl)
            i += 1
            continue

        # ── Regular paragraph ────────────────────────────────────────────────
        para = doc.add_paragraph()
        _set_para_rtl(para, is_rtl)
        para.paragraph_format.space_after = Pt(4)
        _add_inline(para, line, is_rtl, base_size=11)
        i += 1

    doc.save(output_file)


# ── CLI ────────────────────────────────────────────────────────────────────────

def main():
    if not HAS_DOCX:
        print("[MISSING_DEP] python-docx is not installed.")
        print("To install: pip install python-docx")
        sys.exit(2)

    if len(sys.argv) < 2:
        print("Usage: python md_to_word.py <input.md> [output.docx]")
        sys.exit(1)

    input_file = sys.argv[1]

    if not os.path.isfile(input_file):
        print(f"[ERROR] File not found: {input_file}")
        sys.exit(1)

    output_file = sys.argv[2] if len(sys.argv) >= 3 else os.path.splitext(input_file)[0] + ".docx"

    md_to_word(input_file, output_file)
    print(f"[DONE] {output_file}")


if __name__ == "__main__":
    main()
